"""
TaskShot DOCX Builder - WCAG 2.1 AA Compliant Word Document Generation
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Brand Colors as RGB tuples
PRIMARY_PURPLE = (162, 59, 132)    # #a23b84
SECONDARY_PURPLE = (58, 43, 149)   # #3a2b95
ACCENT_PURPLE = (111, 47, 166)     # #6f2fa6


class DocxBuilder:
    """Creates WCAG 2.1 AA compliant Word documents"""

    def __init__(self, settings):
        self.settings = settings
        self._image_width = 4.8   # Screenshot width in inches
        self._image_height = 3.1  # Screenshot height in inches

    def create_document(self, screenshots, task_name):
        """
        Create a Word document from screenshots.

        Args:
            screenshots: List of screenshot data dicts
            task_name: Name of the task/tutorial

        Returns:
            Path to the created document
        """
        doc = Document()

        # Set up document properties
        self._setup_document(doc, task_name)

        # Add header and footer
        self._add_header_footer(doc, task_name)

        # Add title section
        self._add_title_section(doc, task_name)

        # Add screenshot table
        self._add_screenshot_table(doc, screenshots)

        # Save document
        output_path = self._get_output_path(task_name)
        doc.save(str(output_path))

        return output_path

    def _setup_document(self, doc, task_name):
        """Configure document properties and page setup"""
        # Set document properties
        core_props = doc.core_properties
        core_props.title = f"{task_name} - Tutorial"
        core_props.author = "Rocco Catrone of Tech Inclusion Pro"
        core_props.subject = "Accessible Step-by-Step Tutorial"
        core_props.language = "en-US"

        # Set page to landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE

        # Swap width and height for landscape
        new_width = section.page_height
        new_height = section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Set margins
        margin_map = {
            '0.5"': Inches(0.5),
            '0.75"': Inches(0.75),
            '1"': Inches(1)
        }
        margin_size = self.settings.get('margin_size', '0.5"')
        margin = margin_map.get(margin_size, Inches(0.5))

        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def _add_header_footer(self, doc, task_name):
        """Add header and footer to the document"""
        section = doc.sections[0]

        # === Header ===
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Task name on left
        run = header_para.add_run(task_name)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*PRIMARY_PURPLE)

        # Add tab for page number on right
        header_para.add_run("\t\t")

        # Page number (using field code)
        self._add_page_number(header_para)

        # Add accent line under header
        self._add_bottom_border(header_para, PRIMARY_PURPLE)

        # === Footer ===
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Left text
        run = footer_para.add_run("TaskShot - Open Source Accessible Documentation")
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Center accessibility statement (if enabled)
        if self.settings.get('include_footer', True):
            footer_para.add_run("  |  ")
            run = footer_para.add_run("This document created with WCAG 2.1 AA compliance")
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_page_number(self, paragraph):
        """Add page number field to paragraph"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        paragraph.add_run(" of ")

        run2 = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2._r.append(fldChar3)

    def _add_bottom_border(self, paragraph, color):
        """Add a bottom border to a paragraph"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '%02x%02x%02x' % color)

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_title_section(self, doc, task_name):
        """Add title and creation info"""
        # Main title (H1)
        heading = doc.add_heading(f"{task_name} Tutorial", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(*PRIMARY_PURPLE)

        # Creation info
        date_str = datetime.now().strftime("%B %d, %Y")
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = info_para.add_run(f"Created: {date_str} | Made by Rocco Catrone of Tech Inclusion Pro")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Section heading (H2)
        doc.add_paragraph()  # Spacing
        heading2 = doc.add_heading("Step-by-Step Instructions", level=2)
        for run in heading2.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(*SECONDARY_PURPLE)

    def _add_screenshot_table(self, doc, screenshots):
        """Add the main screenshot table"""
        # Create table with header row
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Set column widths to 5 inches each
        step_col_width = 5.0
        screenshot_col_width = 5.0

        # Set column widths
        table.columns[0].width = Inches(step_col_width)
        table.columns[1].width = Inches(screenshot_col_width)

        # Store image dimensions: 4.8 inches wide x 3.1 inches tall
        self._image_width = 4.8
        self._image_height = 3.1

        # Style header row
        header_cells = table.rows[0].cells
        self._style_header_cell(header_cells[0], "Step")
        self._style_header_cell(header_cells[1], "Screenshot & Notes")

        # Mark header row for accessibility
        self._set_table_header_row(table)

        # Add each screenshot
        for i, ss in enumerate(screenshots):
            row = table.add_row()
            row_cells = row.cells

            # Enforce cell widths for this row (5 inches each)
            row_cells[0].width = Inches(5.0)
            row_cells[1].width = Inches(5.0)

            # Alternate row colors
            bg_color = 'FFFFFF' if i % 2 == 0 else 'F5F5F5'
            self._set_cell_background(row_cells[0], bg_color)
            self._set_cell_background(row_cells[1], bg_color)

            # Step column
            step_para = row_cells[0].paragraphs[0]
            run = step_para.add_run(f"Step {i + 1}:\n")
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True

            title_run = step_para.add_run(ss.get('title', 'Untitled'))
            title_run.font.name = 'Arial'
            title_run.font.size = Pt(12)

            # Screenshot column - center the image
            ss_cell = row_cells[1]

            # Add image with alt text (centered)
            self._add_image_with_alt_text(
                ss_cell,
                ss.get('image_path'),
                ss.get('alt_text', 'Screenshot')
            )

            # Add notes if present
            notes = ss.get('notes', '').strip()
            if notes:
                notes_para = ss_cell.add_paragraph()
                notes_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = notes_para.add_run(notes)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.italic = True

        # Apply table borders
        self._apply_table_borders(table)

    def _style_header_cell(self, cell, text):
        """Style a header cell"""
        # Set background color
        self._set_cell_background(cell, '%02x%02x%02x' % SECONDARY_PURPLE)

        # Add text
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    def _set_cell_background(self, cell, color):
        """Set background color for a cell"""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def _set_table_header_row(self, table):
        """Mark the first row as a header row for accessibility"""
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)

        # Get the first row
        first_row = table.rows[0]._tr
        trPr = first_row.get_or_add_trPr()

        # Add table header element
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

    def _add_image_with_alt_text(self, cell, image_path, alt_text):
        """Add an image to a cell with alt text, centered"""
        try:
            para = cell.paragraphs[0]
            # Center the paragraph containing the image
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()

            # Use fixed image dimensions: 4.8 inches wide x 3.1 inches tall
            image_width = getattr(self, '_image_width', 4.8)
            image_height = getattr(self, '_image_height', 3.1)

            # Add image with specified dimensions
            picture = run.add_picture(image_path, width=Inches(image_width), height=Inches(image_height))

            # Set alt text
            inline = picture._inline
            docPr = inline.docPr
            docPr.set('descr', alt_text)
            docPr.set('title', alt_text)

        except Exception as e:
            print(f"Error adding image: {e}")
            # Add placeholder text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[Image: {alt_text}]")
            run.font.italic = True

    def _apply_table_borders(self, table):
        """Apply borders to all table cells"""
        border_color = self.settings.get('border_color', '#6f2fa6')
        # Remove # prefix if present
        if border_color.startswith('#'):
            border_color = border_color[1:]

        tbl = table._tbl
        tblPr = tbl.tblPr

        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')  # 1pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _get_output_path(self, task_name):
        """Generate output file path"""
        # Clean task name for filename
        safe_name = "".join(c for c in task_name if c.isalnum() or c in ' -_').strip()
        safe_name = safe_name.replace(' ', '_')

        date_str = datetime.now().strftime("%Y-%m-%d")
        filename = f"{safe_name}.TaskShot.{date_str}.docx"

        # Save to Desktop
        desktop = Path.home() / "Desktop"
        return desktop / filename
